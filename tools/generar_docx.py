#!/usr/bin/env python3
"""
Tool: generar_docx
Convierte contenido Markdown a un documento .docx con formato profesional.

Activador: se activa automaticamente cuando el agente detecta intencion
de crear un informe de ventas (palabras clave: informe, reporte, report,
documento, word, docx, balance, resumen mensual).

Uso: python generar_docx.py "ruta/de/salida.docx"
     python generar_docx.py "ruta/de/salida.docx" "ruta/al/archivo.md"
"""
import sys
import os
import re
import argparse
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def aplicar_estilo_tabla(table):
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            if row_idx == 0:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1F4E79"/>')
                cell._tc.get_or_add_tcPr().append(shading)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.name = 'Calibri'
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.font.bold = True
                        run.font.size = Pt(10)
            else:
                if row_idx % 2 == 0:
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D6E4F0"/>')
                    cell._tc.get_or_add_tcPr().append(shading)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(9)


def _agregar_parrafo_con_formato(doc, text, style=None):
    """Agrega un párrafo respetando **bold** e _italic_ inline."""
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    patron = re.compile(r'(\*\*.*?\*\*|_.*?_)')
    partes = patron.split(text.replace('⚠️', ''))
    for parte in partes:
        if parte.startswith('**') and parte.endswith('**'):
            run = p.add_run(parte[2:-2])
            run.bold = True
        elif parte.startswith('_') and parte.endswith('_') and len(parte) > 2:
            run = p.add_run(parte[1:-1])
            run.italic = True
        else:
            p.add_run(parte)
    return p


def markdown_a_docx(md_text: str) -> Document:
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for h_level in range(1, 5):
        hs = doc.styles[f'Heading {h_level}']
        hs.font.name = 'Times New Roman'
        hs.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    lines = md_text.split('\n')
    i = 0
    in_table = False
    table_data = []

    while i < len(lines):
        line = lines[i]

        if line.startswith('| ') and line.endswith('|') and '|' in line[2:-2]:
            if not in_table:
                table_data = []
                in_table = True
            cells = [c.strip() for c in line.split('|')[1:-1]]
            table_data.append(cells)
            i += 1
            if i < len(lines) and re.match(r'^[\s\|:\-]+$', lines[i]):
                i += 1
            continue
        elif in_table:
            if table_data:
                rows = len(table_data) - 1
                cols = len(table_data[0])
                table = doc.add_table(rows=rows + 1, cols=cols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for r, row_data in enumerate(table_data):
                    for c, cell_text in enumerate(row_data):
                        table.rows[r].cells[c].text = cell_text
                aplicar_estilo_tabla(table)
                doc.add_paragraph()
            in_table = False
            table_data = []

        text = line.strip()
        if not text:
            if not in_table:
                doc.add_paragraph()
            i += 1
            continue

        heading_match = re.match(r'^(#{1,6})\s+(.*)', text)
        if heading_match:
            hashes = heading_match.group(1)
            title = heading_match.group(2).strip()
            level = min(len(hashes), 4)  # Word soporta Heading 1-4 nativo
            h = doc.add_heading(title, level=level)
            for run in h.runs:
                run.font.name = 'Times New Roman'
                if level == 1:
                    run.font.size = Pt(28)
                    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                elif level == 2:
                    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
                elif level == 3:
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
                else:
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
            i += 1
            continue

        if re.match(r'^-\s', text):
            text_clean = re.sub(r'^-\s+', '', text)
            _agregar_parrafo_con_formato(doc, text_clean, style='List Bullet')
            i += 1
            continue

        if re.match(r'^\d+\.\s', text):
            text_clean = re.sub(r'^\d+\.\s+', '', text)
            _agregar_parrafo_con_formato(doc, text_clean, style='List Number')
            i += 1
            continue

        if text.startswith('---'):
            doc.add_paragraph('_' * 60)
            i += 1
            continue

        if re.match(r'^\*\*', text):
            # Linea que empieza con bold — puede tener mezcla inline
            _agregar_parrafo_con_formato(doc, text)
            i += 1
            continue

        img_match = re.match(r'^!\[(.*?)\]\((.*?)\)', text)
        if img_match:
            alt_text = img_match.group(1)
            img_path = img_match.group(2).strip()
            if os.path.exists(img_path):
                try:
                    doc.add_picture(img_path, width=Inches(5.5))
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = 1  # CENTER
                    p = doc.add_paragraph()
                    run = p.add_run(alt_text)
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
                    p.alignment = 1  # CENTER
                except Exception:
                    doc.add_paragraph(f'[Grafico: {alt_text}]')
            else:
                doc.add_paragraph(f'[Grafico no disponible: {alt_text}]')
            i += 1
            continue

        _agregar_parrafo_con_formato(doc, text)
        i += 1

    if in_table and table_data:
        rows = len(table_data) - 1
        cols = len(table_data[0])
        table = doc.add_table(rows=rows + 1, cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for r, row_data in enumerate(table_data):
            for c, cell_text in enumerate(row_data):
                table.rows[r].cells[c].text = cell_text
        aplicar_estilo_tabla(table)

    return doc


def main():
    parser = argparse.ArgumentParser(description='Convierte Markdown a DOCX')
    default_out = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'reports', 'informe_ventas.docx') if 'MD_CONTENT' in os.environ else 'informe_ventas.docx'
    parser.add_argument('salida', nargs='?', default=default_out, help='Ruta del archivo .docx de salida')
    parser.add_argument('entrada', nargs='?', help='Ruta al archivo .md')
    parser.add_argument('--inline', action='store_true', help='Usar entrada como contenido inline')
    args = parser.parse_args()

    ruta_salida = args.salida or 'informe_ventas.docx'

    md_content = os.environ.get('MD_CONTENT')
    if md_content:
        pass
    elif args.entrada and os.path.exists(args.entrada):
        with open(args.entrada, 'r', encoding='utf-8') as f:
            md_content = f.read()
    else:
        print(json.dumps({'success': False, 'error': 'No se pudo obtener el contenido. Usa MD_CONTENT env o pasa un archivo.'}))
        sys.exit(1)

    doc = markdown_a_docx(md_content)
    doc.save(args.salida)
    print(json.dumps({'success': True, 'archivo': os.path.abspath(args.salida)}))


if __name__ == '__main__':
    import json
    main()
